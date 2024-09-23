from docx import Document

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('Index of Python Record Programs', 0)

# Data for the index
index_data = [
    ("Program #1.1", "Simple Calculator", "1"),
    ("Program #1.2", "Electric Power Distribution Bill Calculation", "1"),
    ("Program #1.3", "Number Pyramid", "2"),
    ("Program #1.4", "Sum and Count of Integers Divisible by 7", "2"),
    ("Program #1.5", "Recursive Sum Calculation", "2"),
    ("Program #1.6", "Reverse and Add Until Palindrome", "2"),
    ("Program #1.7", "String Operations", "3"),
    ("Program #1.8", "Factorial with Memoization", "3"),
    ("Program #1.9", "Set Operations", "3"),
    ("Program #1.10", "Student Details Management", "3"),
    ("Program #1.11", "File Copy", "4"),
    ("Program #1.12", "Directory Operations Using OS Module", "4"),
    ("Program #1.13", "Banking Application Using Inheritance", "5"),
    ("Program #2.1", "MySQL Connection with Python", "5"),
    ("Program #2.2", "SQLite Connection with Python", "6"),
    ("Program #2.3", "MySQL CRUD Operations with Python", "6"),
    ("Program #3.1", "CGI-Based Login Form", "7"),
    ("Program #3.2", "MCA Admission Registration Form", "7"),
    ("Program #3.3", "MySQL CRUD with CGI", "8"),
    ("Program #4.1.1", "Numpy Ones Array", "8"),
    ("Program #4.1.2", "Remove Non-Numeric Rows in Numpy", "9"),
    ("Program #4.1.3", "Remove Single-Dimensional Entries", "9"),
    ("Program #4.1.4", "Check Values in Numpy Array", "9"),
    ("Program #4.1.5", "2D Diagonals of 3D Numpy Array", "10"),
    ("Program #4.1.6", "Sort Numpy Array by Axis", "10"),
    ("Program #4.1.7", "Sort Structured Array", "11"),
    ("Program #4.1.8", "Sort Complex Array", "11"),
    ("Program #4.1.9", "Sort by Column in Numpy", "12"),
    ("Program #4.1.10", "Sum of Diagonal Elements", "12"),
    ("Program #4.1.11", "Matrix Multiplication", "13"),
    ("Program #4.1.12", "Matrix Multiplication of Complex Numbers", "13"),
    ("Program #4.1.13", "Inner, Outer, and Cross Products", "14"),
    ("Program #4.1.14", "Covariance Matrix", "14"),
    ("Program #4.1.15", "Covariance to Correlation Matrix", "15"),
    ("Program #4.1.16", "Histogram Calculation", "15"),
    ("Program #4.1.17", "Cross-Correlation Calculation", "16"),
    ("Program #4.1.18", "Mean, Standard Deviation, and Variance", "16"),
    ("Program #4.1.19", "Calculate 80th Percentile", "17"),
    ("Program #4.2.1", "Pandas Series Arithmetic", "17"),
    ("Program #4.2.2", "Convert Dictionary to Pandas Series", "18"),
    ("Program #4.2.3", "Convert DataFrame Column to Series", "18"),
    ("Program #4.2.4", "Convert Series of Lists to Series", "18"),
    ("Program #4.2.5", "Create Subset of Series Based on Condition", "19"),
    ("Program #4.2.6", "Find Uncommon Items Between Series", "19"),
    ("Program #4.2.7", "Frequency Count of Series Values", "19"),
    ("Program #4.2.8", "Filter Words with Vowels", "20"),
    ("Program #4.2.9", "Index of Min and Max Values in Series", "20"),
    ("Program #4.2.10", "Get First 3 Rows of DataFrame", "21"),
    ("Program #4.2.11", "Select Columns from DataFrame", "21"),
    ("Program #4.2.12", "Count Rows and Columns in DataFrame", "21"),
    ("Program #4.2.13", "Add Row to DataFrame", "22"),
    ("Program #4.2.14", "Write DataFrame to CSV with Tab Separator", "22"),
    ("Program #4.2.15", "Replace NaN Values and Drop Rows", "23"),
    ("Program #4.2.16", "Shuffle DataFrame Rows", "23"),
    ("Program #4.2.17", "Find Row with Max Column Value", "23"),
    ("Program #4.2.18", "Check for Column in DataFrame", "24"),
    ("Program #4.2.19", "Append Data to Empty DataFrame", "24"),
    ("Program #4.2.20", "Convert Continuous Values to Categorical", "25"),
    ("Program #4.2.21", "Create DataFrame with Random and Missing Values", "25"),
    ("Program #4.2.22", "Join DataFrames Along Rows", "26"),
    ("Program #4.2.23", "Join DataFrames Along Columns", "26")
]

# Add table headers
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Program Number'
hdr_cells[1].text = 'Name'
hdr_cells[2].text = 'Page Number'

# Populate the table with data
for program_num, name, page_num in index_data:
    row_cells = table.add_row().cells
    row_cells[0].text = program_num
    row_cells[1].text = name
    row_cells[2].text = page_num

# Save the document
doc_path = 'Python_Record_Index.docx'
doc.save(doc_path)

doc_path
