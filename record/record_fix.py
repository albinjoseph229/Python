from docx import Document
from docx.shared import Inches

# Open the user's Word document
doc_path = 'd:/MCA_PROGRAMS/3rdSEM/Python/record/saalim_python_final.docx'
doc = Document(doc_path)

# Adjust the footer and indentation
for paragraph in doc.paragraphs:
    if paragraph.alignment is None:  # Adjust non-aligned paragraphs
        paragraph.paragraph_format.left_indent = Inches(0.5)  # Indent to 0.5 inches

# Update footers
for section in doc.sections:
    footer = section.footer
    for paragraph in footer.paragraphs:
        paragraph.text = "Corrected Footer: Data Analytics using Python"
        paragraph.alignment = 1  # Center the footer text

# Save the modified document
updated_doc_path = 'd:/MCA_PROGRAMS/3rdSEM/Python/record/saalim_python_final_corrected_with_footer.docx'
doc.save(updated_doc_path)

print("File saved as:", updated_doc_path)
